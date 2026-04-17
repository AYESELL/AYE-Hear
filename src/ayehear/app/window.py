from __future__ import annotations

import logging
import threading
from datetime import datetime
from typing import TYPE_CHECKING

import numpy as np

from PySide6.QtCore import Qt, QTimer, Signal
from PySide6.QtWidgets import (
    QAbstractItemView,
    QComboBox,
    QFormLayout,
    QGroupBox,
    QHBoxLayout,
    QLabel,
    QLineEdit,
    QListWidget,
    QListWidgetItem,
    QMainWindow,
    QMessageBox,
    QPushButton,
    QPlainTextEdit,
    QSplitter,
    QSpinBox,
    QVBoxLayout,
    QWidget,
)

from ayehear.app.enrollment_dialog import EnrollmentDialog
from ayehear.app.mic_level_widget import MicLevelWidget
from ayehear.app.system_readiness import ReadinessChecker, SystemReadinessWidget
from ayehear.models.meeting import MeetingSession, Participant
from ayehear.models.runtime import RuntimeConfig
from ayehear.services.audio_capture import (
    AudioCaptureProfile,
    AudioCaptureService,
    AudioSegment,
    enumerate_input_devices,
)
from ayehear.services.protocol_engine import ProtocolEngine
from ayehear.services.speaker_manager import SpeakerManager
from ayehear.services.transcription import TranscriptionService
from ayehear.storage.database import DatabaseBootstrap, DatabaseConfig, load_runtime_dsn
from ayehear.storage.repositories import (
    ProtocolSnapshotRepository,
    SpeakerProfileRepository,
)

if TYPE_CHECKING:
    from sqlalchemy.orm import Session
    from ayehear.storage.repositories import (
        MeetingRepository,
        ParticipantRepository,
        TranscriptSegmentRepository,
        ProtocolSnapshotRepository,
    )

logger = logging.getLogger(__name__)

# Prefix used to identify degraded-state placeholder text in the protocol view
# so exports and tests can detect when the panel shows no real content.
_PROTOCOL_DEGRADED_PREFIX = "[DEGRADED]"


class MainWindow(QMainWindow):
    transcript_line_ready = Signal(str)

    def __init__(
        self,
        runtime_config: RuntimeConfig,
        db_session: "Session | None" = None,
        meeting_repo: "MeetingRepository | None" = None,
        participant_repo: "ParticipantRepository | None" = None,
        transcript_repo: "TranscriptSegmentRepository | None" = None,
        snapshot_repo: "ProtocolSnapshotRepository | None" = None,
        speaker_manager: "SpeakerManager | None" = None,
    ) -> None:
        super().__init__()
        self.runtime_config = runtime_config
        self._db_session = db_session
        self._meeting_repo = meeting_repo
        self._participant_repo = participant_repo
        self._transcript_repo = transcript_repo
        self._snapshot_repo = snapshot_repo
        self._active_meeting_id: str | None = None
        self._session: MeetingSession | None = None
        self._audio_capture_service: AudioCaptureService | None = None
        self._transcription_service = TranscriptionService(
            profile=self.runtime_config.models.whisper_profile,
            language="de",
            transcript_repo=transcript_repo,
        )
        # ADR-0003: speaker identification pipeline
        self._speaker_manager: SpeakerManager = speaker_manager or SpeakerManager()
        self._enrolled_speakers: dict[str, str] = {}  # participant_id -> profile_id
        # Maps list-item UUID (UserRole) -> DB Participant.id (HEAR-084 AC1/AC3)
        self._participant_id_map: dict[str, str] = {}
        # HEAR-085 AC2: protocol engine for structured draft generation (ADR-0005)
        self._protocol_engine = ProtocolEngine(
            snapshot_repo=snapshot_repo,
            transcript_repo=transcript_repo,
        )
        # HEAR-087: system readiness checker + widget (built in _build_setup_panel)
        self._readiness_checker = ReadinessChecker()
        self._readiness_widget: SystemReadinessWidget | None = None
        self._audio_buffer_lock = threading.Lock()
        self._pending_audio_chunks: list[np.ndarray] = []
        self._pending_start_ms: int | None = None
        self._pending_end_ms: int = 0
        self._pending_duration_ms: int = 0
        self._asr_warned_no_text = False

        self.setWindowTitle("AYE Hear")
        self.resize(1440, 900)

        central = QWidget(self)
        layout = QVBoxLayout(central)
        layout.setContentsMargins(16, 16, 16, 16)
        layout.setSpacing(12)

        header = QLabel("AYE Hear Workspace")
        header.setObjectName("pageTitle")
        header.setStyleSheet("font-size: 24px; font-weight: 700;")
        layout.addWidget(header)

        splitter = QSplitter(Qt.Orientation.Horizontal)
        splitter.addWidget(self._build_setup_panel())
        splitter.addWidget(self._build_transcript_panel())
        splitter.addWidget(self._build_protocol_panel())
        splitter.setSizes([320, 500, 500])
        layout.addWidget(splitter)

        self.setCentralWidget(central)

        # Auto-refresh review queue every 10 s while a meeting is active
        self._refresh_timer = QTimer(self)
        self._refresh_timer.setInterval(10_000)
        self._refresh_timer.timeout.connect(self._refresh_review_queue)
        # HEAR-085 AC2: also rebuild structured protocol draft on each refresh tick
        self._refresh_timer.timeout.connect(self._rebuild_protocol_from_persistence)

        self._asr_timer = QTimer(self)
        self._asr_timer.setInterval(1000)
        self._asr_timer.timeout.connect(self._process_pending_audio)
        self.transcript_line_ready.connect(self.append_transcript_line)

    # ------------------------------------------------------------------
    # Panel builders
    # ------------------------------------------------------------------

    # ------------------------------------------------------------------
    # Panel builders
    # ------------------------------------------------------------------
    def _build_setup_panel(self) -> QWidget:
        panel = QWidget()
        layout = QVBoxLayout(panel)
        layout.setSpacing(12)

        meeting_box = QGroupBox("Meeting Setup")
        form = QFormLayout(meeting_box)

        self._meeting_title = QLineEdit()
        form.addRow("Meeting Title", self._meeting_title)

        self._meeting_type = QComboBox()
        self._meeting_type.addItems(self.runtime_config.protocol.meeting_modes)
        form.addRow("Meeting Type", self._meeting_type)

        self._participant_count = QSpinBox()
        self._participant_count.setRange(1, 30)
        self._participant_count.setValue(2)
        form.addRow("Participants", self._participant_count)

        self._naming_template = QComboBox()
        self._naming_template.addItems([
            "Herr/Frau + Last Name + Company",
            "First Name + Last Name + Company",
        ])
        form.addRow("Participant Template", self._naming_template)

        # HEAR-039: real Windows device selector
        self._audio_device = QComboBox()
        self._audio_device.addItem("Windows default microphone", userData=None)
        self._populate_audio_devices()
        form.addRow("Audio Input", self._audio_device)

        # HEAR-093: Protocol language selection (DE/EN/FR)
        self._protocol_language = QComboBox()
        self._protocol_language.addItems(
            self.runtime_config.protocol.protocol_language_options
        )
        default_lang = self.runtime_config.protocol.protocol_language
        idx = self._protocol_language.findText(default_lang)
        if idx >= 0:
            self._protocol_language.setCurrentIndex(idx)
        self._protocol_language.currentTextChanged.connect(self._on_protocol_language_changed)
        form.addRow("Protocol Language", self._protocol_language)

        layout.addWidget(meeting_box)

        speakers_box = QGroupBox("Speaker Enrollment")
        speakers_layout = QVBoxLayout(speakers_box)
        speakers_layout.addWidget(QLabel(
            "Doppelklick zum Bearbeiten. Format: Name | Organisation | Status"
        ))

        self._speakers_list = QListWidget()
        self._speakers_list.setEditTriggers(
            QAbstractItemView.EditTrigger.DoubleClicked
            | QAbstractItemView.EditTrigger.EditKeyPressed
        )
        for entry in [
            "Frau Schneider | AYE | pending enrollment",
            "Max Weber | Customer GmbH | pending enrollment",
        ]:
            import uuid as _uuid
            item = QListWidgetItem(entry)
            item.setFlags(item.flags() | Qt.ItemFlag.ItemIsEditable)
            item.setData(Qt.ItemDataRole.UserRole, str(_uuid.uuid4()))
            self._speakers_list.addItem(item)
        speakers_layout.addWidget(self._speakers_list)
        # HEAR-040: update status label after any user-committed inline edit
        self._speakers_list.itemChanged.connect(self._on_speaker_item_changed)

        speaker_btn_row = QHBoxLayout()
        add_btn = QPushButton("+ Add")
        add_btn.setToolTip("Neuen Sprecher hinzufügen")
        add_btn.clicked.connect(self._add_speaker)
        edit_btn = QPushButton("Edit")
        edit_btn.setToolTip("Ausgewählten Sprecher bearbeiten")
        edit_btn.clicked.connect(self._edit_speaker)
        remove_btn = QPushButton("Remove")
        remove_btn.setToolTip("Ausgewählten Sprecher entfernen")
        remove_btn.clicked.connect(self._remove_speaker)
        speaker_btn_row.addWidget(add_btn)
        speaker_btn_row.addWidget(edit_btn)
        speaker_btn_row.addWidget(remove_btn)
        speakers_layout.addLayout(speaker_btn_row)

        # HEAR-040: visible feedback label for speaker actions
        self._speaker_status = QLabel("")
        self._speaker_status.setStyleSheet("color: #1a7a1a; font-weight: 600;")
        speakers_layout.addWidget(self._speaker_status)

        apply_template_btn = QPushButton("Apply Participant Template")
        apply_template_btn.clicked.connect(self._apply_participant_template)
        speakers_layout.addWidget(apply_template_btn)

        start_enrollment_btn = QPushButton("Start Enrollment")
        start_enrollment_btn.clicked.connect(self._start_enrollment)
        speakers_layout.addWidget(start_enrollment_btn)

        layout.addWidget(speakers_box)

        # HEAR-087: system readiness indicators
        self._readiness_widget = SystemReadinessWidget()
        self._readiness_widget.set_refresh_callback(self._refresh_readiness)
        layout.addWidget(self._readiness_widget)

        # HEAR-041: meeting status indicator
        self._meeting_status_label = QLabel("\u26aa Kein aktives Meeting")
        self._meeting_status_label.setStyleSheet("font-weight: 600; color: #888;")
        layout.addWidget(self._meeting_status_label)
        # HEAR-044: live mic state + level meter
        self._mic_level_widget = MicLevelWidget()
        layout.addWidget(self._mic_level_widget)
        controls = QHBoxLayout()
        self._start_meeting_btn = QPushButton("Start Meeting")
        self._start_meeting_btn.clicked.connect(self._start_meeting)
        self._stop_meeting_btn = QPushButton("Stop Meeting")
        self._stop_meeting_btn.clicked.connect(self._stop_meeting)
        self._stop_meeting_btn.setEnabled(False)
        show_state_btn = QPushButton("Show Current State")
        show_state_btn.clicked.connect(self._show_current_state)
        controls.addWidget(self._start_meeting_btn)
        controls.addWidget(self._stop_meeting_btn)
        controls.addWidget(show_state_btn)
        layout.addLayout(controls)
        layout.addStretch(1)
        return panel

    # ------------------------------------------------------------------
    # HEAR-039: audio device helpers
    # ------------------------------------------------------------------

    def _populate_audio_devices(self) -> None:
        """Fill the Audio Input dropdown with available Windows capture devices."""
        devices = enumerate_input_devices()
        if devices:
            self._audio_device.clear()
            for idx, name in devices:
                self._audio_device.addItem(name, userData=idx)
        # If no devices found the fallback item set during construction remains.
        # HEAR-087: refresh readiness after devices are enumerated
        QTimer.singleShot(0, self._refresh_readiness)

    def _selected_audio_profile(self) -> AudioCaptureProfile:
        """Return an AudioCaptureProfile for the selected input device (HEAR-039).

        device_index is the sounddevice index (int) or None for WASAPI default.
        Use this when starting AudioCaptureService once the audio pipeline is
        integrated (ADR-0004).
        """
        device_index: int | None = self._audio_device.currentData()
        return AudioCaptureProfile(device_index=device_index)

    # ------------------------------------------------------------------
    # HEAR-087: system readiness
    # ------------------------------------------------------------------

    def _refresh_readiness(self) -> None:
        """Evaluate all component readiness states and update the indicator widget.
        
        This method now attempts to reload the persistence layer each time
        (in case the DSN file was created after app startup), then re-evaluates
        all component readiness states.
        """
        # Try to reload persistence layer if not already connected
        if self._meeting_repo is None:
            self._reload_persistence_layer()
        
        if self._readiness_widget is None:
            return
        components, aggregate = self._readiness_checker.check_all(
            meeting_repo=self._meeting_repo,
            participant_repo=self._participant_repo,
            transcript_repo=self._transcript_repo,
            snapshot_repo=self._snapshot_repo,
            speaker_profile_repo=self._speaker_manager._profiles,
            runtime_config=self.runtime_config,
        )
        self._readiness_widget.update_status(components, aggregate)

    def _reload_persistence_layer(self) -> None:
        """Attempt to load or reload the persistence layer from DSN.
        
        Called by _refresh_readiness() to support dynamic DSN discovery
        when pg.dsn is written after app startup (e.g., post-provisioning).
        """
        try:
            dsn = load_runtime_dsn()
            if dsn:
                logger.info("Loading runtime DSN for persistence layer reload.")
                bootstrap = DatabaseBootstrap(DatabaseConfig(dsn=dsn))
                bootstrap.bootstrap()
                new_session = bootstrap.session()
                
                # Close old session if it exists
                if self._db_session is not None:
                    self._db_session.close()
                
                self._db_session = new_session
                
                # Reinitialize all repositories
                from ayehear.storage.repositories import (
                    MeetingRepository,
                    ParticipantRepository,
                    TranscriptSegmentRepository,
                )
                self._meeting_repo = MeetingRepository(new_session)
                self._participant_repo = ParticipantRepository(new_session)
                self._transcript_repo = TranscriptSegmentRepository(new_session)
                self._snapshot_repo = ProtocolSnapshotRepository(new_session)
                
                # Reinitialize speaker manager with new session
                self._speaker_manager = SpeakerManager(
                    profile_repo=SpeakerProfileRepository(new_session),
                    participant_repo=self._participant_repo,
                )
                
                # Update services that depend on repos
                self._transcription_service._transcript_repo = self._transcript_repo
                self._protocol_engine._snapshot_repo = self._snapshot_repo
                self._protocol_engine._transcript_repo = self._transcript_repo
                
                logger.info("Persistence layer reloaded successfully.")
        except Exception as exc:
            logger.error("Failed to reload persistence layer: %s", exc)


    # ------------------------------------------------------------------
    # Speaker management (HEAR-036 / HEAR-040)
    # ------------------------------------------------------------------

    # HEAR-093: Protocol language change handler
    def _on_protocol_language_changed(self, language: str) -> None:
        """Propagate selected protocol language to ProtocolEngine."""
        self._protocol_engine._language = language
        logger.debug("Protocol language set to: %s", language)

    def _on_speaker_item_changed(self, item: QListWidgetItem) -> None:
        """Update feedback label after user commits an inline edit (HEAR-040)."""
        self._set_speaker_status(f"Gespeichert: {item.text()[:40]}")

    def _set_speaker_status(self, message: str) -> None:
        """Show a short status message below the speaker list."""
        self._speaker_status.setText(message)

    def _add_speaker(self) -> None:
        """Add a new editable speaker entry to the list."""
        import uuid as _uuid
        item = QListWidgetItem("Name | Organisation | pending enrollment")
        item.setFlags(item.flags() | Qt.ItemFlag.ItemIsEditable)
        item.setData(Qt.ItemDataRole.UserRole, str(_uuid.uuid4()))
        self._speakers_list.addItem(item)
        self._speakers_list.setCurrentItem(item)
        self._speakers_list.editItem(item)
        self._set_speaker_status("Neuer Sprecher hinzugef\u00fcgt \u2014 Namen direkt editieren.")

    def _edit_speaker(self) -> None:
        """Inline-Edit des ausgewählten Sprecher-Eintrags."""
        item = self._speakers_list.currentItem()
        if item is None:
            QMessageBox.information(self, "Edit Speaker", "Bitte einen Sprecher ausw\u00e4hlen.")
            self._set_speaker_status("Kein Sprecher ausgew\u00e4hlt.")
            return
        self._speakers_list.editItem(item)
        self._set_speaker_status(f"Bearbeite: {item.text()[:40]}")

    def _remove_speaker(self) -> None:
        """Entfernt den ausgewählten Sprecher nach Bestätigung."""
        item = self._speakers_list.currentItem()
        if item is None:
            QMessageBox.information(self, "Remove Speaker", "Bitte einen Sprecher ausw\u00e4hlen.")
            self._set_speaker_status("Kein Sprecher ausgew\u00e4hlt.")
            return
        answer = QMessageBox.question(
            self,
            "Sprecher entfernen",
            f"'{item.text()}' entfernen?",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
        )
        if answer == QMessageBox.StandardButton.Yes:
            name = item.text()
            self._speakers_list.takeItem(self._speakers_list.row(item))
            self._set_speaker_status(f"Sprecher entfernt: {name[:40]}")

    def _get_speaker_texts(self) -> list[str]:
        """Gibt alle nicht-leeren Sprecher-Texte zurück."""
        result = []
        for i in range(self._speakers_list.count()):
            text = self._speakers_list.item(i).text().strip()
            if text:
                result.append(text)
        return result

    # ------------------------------------------------------------------
    # Setup action handlers (HEAR-037)
    # ------------------------------------------------------------------

    def _apply_participant_template(self) -> None:
        """Generiert Platzhalter-Sprecher aus Anzahl und Namens-Template."""
        import uuid as _uuid
        count = self._participant_count.value()
        use_salutation = self._naming_template.currentText().startswith("Herr/Frau")
        self._speakers_list.clear()
        for i in range(1, count + 1):
            if use_salutation:
                text = f"Herr/Frau Teilnehmer_{i} | Organisation | pending enrollment"
            else:
                text = f"Vorname_{i} Nachname_{i} | Organisation | pending enrollment"
            item = QListWidgetItem(text)
            item.setFlags(item.flags() | Qt.ItemFlag.ItemIsEditable)
            item.setData(Qt.ItemDataRole.UserRole, str(_uuid.uuid4()))
            self._speakers_list.addItem(item)

    def _start_enrollment(self) -> None:
        """Open the real voice enrollment dialog (HEAR-074 / ADR-0003 Stage 1).

        Collects pending speakers from the list, opens EnrollmentDialog for
        microphone-based recording, and updates list item status after each
        successful enrollment.  The dialog guides the user with on-screen
        instructions and shows status transitions (pending → recording →
        enrolled / failed).

        Speaker items are identified by a stable UUID stored in
        ``Qt.ItemDataRole.UserRole`` so that display-name changes do not break
        enrollment linkage (HEAR-079).
        """
        import uuid as _uuid
        pending: list[tuple[str, str, str]] = []
        for i in range(self._speakers_list.count()):
            item = self._speakers_list.item(i)
            raw = item.text().strip()
            if not raw:
                continue
            name, org, status = self._parse_speaker_raw(raw)
            if "enrolled" not in status.lower():
                participant_id = item.data(Qt.ItemDataRole.UserRole)
                if not participant_id:
                    participant_id = str(_uuid.uuid4())
                    item.setData(Qt.ItemDataRole.UserRole, participant_id)
                pending.append((name, org, participant_id))

        # No speakers at all → warning
        if self._speakers_list.count() == 0:
            QMessageBox.warning(
                self,
                "Enrollment",
                "Bitte mindestens einen Sprecher hinzufügen.",
            )
            return

        # All speakers already enrolled → informational hint
        if not pending:
            QMessageBox.information(
                self,
                "Enrollment",
                "Alle Sprecher sind bereits enrolliert.",
            )
            return

        dlg = EnrollmentDialog(
            pending_speakers=pending,
            speaker_manager=self._speaker_manager,
            parent=self,
        )
        from PySide6.QtWidgets import QDialog
        accepted = dlg.exec() == QDialog.DialogCode.Accepted
        enrolled = dlg.get_enrolled_results() if accepted else {}
        # Update list items by stable participant_id (HEAR-079)
        pending_ids = {pid for _, _, pid in pending}
        for i in range(self._speakers_list.count()):
            item = self._speakers_list.item(i)
            participant_id = item.data(Qt.ItemDataRole.UserRole)
            if not participant_id:
                continue
            name, org, _ = self._parse_speaker_raw(item.text().strip())
            if participant_id in enrolled:
                profile_id = enrolled[participant_id]
                item.setText(f"{name} | {org} | enrolled (id: {profile_id[:8]})")
                self._enrolled_speakers[participant_id] = profile_id
                # HEAR-084 AC3: persist participant-to-profile linkage when repo available
                db_participant_id = self._participant_id_map.get(participant_id)
                if db_participant_id and self._participant_repo is not None:
                    try:
                        self._participant_repo.mark_enrolled(db_participant_id, profile_id)
                        logger.info(
                            "Enrollment persisted: participant=%s profile=%s",
                            db_participant_id, profile_id,
                        )
                    except Exception as exc:
                        logger.error("Failed to persist enrollment linkage: %s", exc)
            elif accepted and participant_id in pending_ids:
                # Dialog was accepted but this speaker was not recorded
                item.setText(f"{name} | {org} | enrollment failed")
        enrolled_count = len(enrolled)
        self._set_speaker_status(
            f"Enrollment abgeschlossen: {enrolled_count}/{len(pending)} Sprecher registriert."
        )

    @staticmethod
    def _parse_speaker_raw(raw: str) -> tuple[str, str, str]:
        """Parse a speaker list entry into (name, org, status) tuple."""
        parts = [p.strip() for p in raw.split("|")]
        name = parts[0] if parts else raw
        org = parts[1] if len(parts) > 1 else "Organisation"
        status = parts[2] if len(parts) > 2 else ""
        return name, org, status

    def _start_meeting(self) -> None:
        """Validiert das Setup und startet eine Meeting-Session."""
        title = self._meeting_title.text().strip()
        if not title:
            QMessageBox.warning(self, "Start Meeting", "Bitte einen Meeting-Titel eingeben.")
            self._meeting_title.setFocus()
            return

        speakers = self._get_speaker_texts()
        if not speakers:
            QMessageBox.warning(self, "Start Meeting", "Bitte mindestens einen Sprecher hinzufügen.")
            return

        import uuid

        participants: list[Participant] = []
        for s in speakers:
            parts = [p.strip() for p in s.split("|")]
            display = parts[0] if parts else s
            org = parts[1] if len(parts) > 1 else ""
            name_parts = display.split()
            last = name_parts[-1] if name_parts else display
            first: str | None = " ".join(name_parts[:-1]) if len(name_parts) > 1 else None
            participants.append(
                Participant(first_name=first, last_name=last, organization=org, role="participant")
            )

        device_label = self._audio_device.currentText()

        self._session = MeetingSession(
            title=title,
            mode=self._meeting_type.currentText(),
            meeting_type=self._meeting_type.currentText(),
            participants=participants,
            started_at=datetime.now(),
        )

        meeting_id = str(uuid.uuid4())

        # HEAR-084 AC1: persist meeting + participants to DB when repos available
        self._participant_id_map = {}
        if self._meeting_repo is not None:
            try:
                db_meeting = self._meeting_repo.create(
                    title=title,
                    meeting_type=self._meeting_type.currentText(),
                    mode=self._meeting_type.currentText(),
                )
                self._meeting_repo.start(db_meeting.id)
                meeting_id = db_meeting.id
                logger.info("Meeting persisted to DB: %s", meeting_id)
            except Exception as exc:
                logger.error("Failed to persist meeting to DB: %s — using in-memory id", exc)

        if self._participant_repo is not None and self._session is not None:
            # Build participant_id_map: list-item UUID -> DB Participant.id
            for i in range(self._speakers_list.count()):
                item = self._speakers_list.item(i)
                list_uuid = item.data(Qt.ItemDataRole.UserRole)
                raw = item.text().strip()
                if not raw or not list_uuid:
                    continue
                p_name, p_org, _ = self._parse_speaker_raw(raw)
                name_parts = p_name.split()
                first_name = " ".join(name_parts[:-1]) if len(name_parts) > 1 else None
                last_name = name_parts[-1] if name_parts else p_name
                try:
                    db_participant = self._participant_repo.add(
                        meeting_id=meeting_id,
                        display_name=p_name,
                        first_name=first_name,
                        last_name=last_name,
                        organization=p_org,
                    )
                    self._participant_id_map[list_uuid] = db_participant.id
                    logger.debug("Participant persisted: %s -> %s", p_name, db_participant.id)
                except Exception as exc:
                    logger.error("Failed to persist participant '%s': %s", p_name, exc)

        known_speakers = [p.display_name for p in self._session.participants]
        # ADR-0003 Stage 0: register participant names for constrained intro matching
        self._speaker_manager.register_meeting_participants(known_speakers)
        self.set_active_meeting(meeting_id, known_speakers=known_speakers)
        audio_status = self._start_audio_pipeline()

        # HEAR-041: transcript + protocol reflect meeting start
        self.append_transcript_line(
            f"[00:00] Meeting '{title}' gestartet — {len(speakers)} Teilnehmer "
            f"| Audio: {device_label}."
        )
        self.append_transcript_line(f"[00:00] System: {audio_status}")

        # HEAR-085 AC1/AC2: protocol panel shows structured state, never transcript mirror
        if self._snapshot_repo is not None:
            self._protocol_view.setPlainText(
                f"# {title}\n"
                f"Typ: {self._meeting_type.currentText()} | Sprecher: {len(speakers)}\n\n"
                "Aufnahme läuft — strukturierter Protokollentwurf erscheint nach der ersten "
                "Snapshot-Generierung (ca. 10 s nach dem ersten Transkript).\n"
            )
        else:
            self._protocol_view.setPlainText(
                "[DEGRADED] Protokollentwurf nicht verfügbar — "
                "Persistenz-Backend nicht verbunden.\n\n"
                "Live-Transkript siehe Transkript-Panel."
            )

        # HEAR-041: visible session state
        self._meeting_status_label.setText(f"\U0001f7e2 Meeting aktiv: {title}")
        self._meeting_status_label.setStyleSheet("font-weight: 700; color: #1a7a1a;")
        self._start_meeting_btn.setEnabled(False)
        self._stop_meeting_btn.setEnabled(True)
        # HEAR-075: enable export while meeting is active
        self._export_btn.setEnabled(True)
        self._export_path_label.setText("")

        QMessageBox.information(self, "Meeting gestartet", f"Meeting '{title}' ist jetzt aktiv.")
        # HEAR-087: refresh readiness state after meeting start
        self._refresh_readiness()

    def _stop_meeting(self) -> None:
        """Beendet die aktive Meeting-Session (HEAR-041)."""
        self._transcribe_pending_buffer(force=True)
        self._stop_audio_pipeline()

        # HEAR-070: export artifacts before clearing session state
        meeting_id = self._active_meeting_id
        title = (self._session.title if self._session is not None else None) or ""
        self._export_meeting_artifacts(meeting_id, title)

        # HEAR-084 AC4: mark meeting as ended in DB
        if meeting_id and self._meeting_repo is not None:
            try:
                self._meeting_repo.end(meeting_id)
                logger.info("Meeting ended in DB: %s", meeting_id)
            except Exception as exc:
                logger.error("Failed to end meeting in DB: %s", exc)

        self.stop_active_meeting()
        self._session = None
        self._meeting_status_label.setText("\u26aa Kein aktives Meeting")
        self._meeting_status_label.setStyleSheet("font-weight: 600; color: #888;")
        self._start_meeting_btn.setEnabled(True)
        self._stop_meeting_btn.setEnabled(False)
        # HEAR-075: keep export accessible after recording stops
        # (export button stays enabled so user can export the final protocol)
        self.append_transcript_line("[--:--] Meeting beendet.")

    def _show_current_state(self) -> None:
        """Zeigt einen Dialog mit dem aktuellen Setup- und Session-Status (HEAR-041)."""
        title = self._meeting_title.text().strip() or "(nicht gesetzt)"
        mode = self._meeting_type.currentText()
        device = self._audio_device.currentText()
        speakers = self._get_speaker_texts()

        lines = [
            f"Meeting-Titel:  {title}",
            f"Meeting-Typ:    {mode}",
            f"Audio-Ger\u00e4t:   {device}",
            f"Sprecher ({len(speakers)}):",
        ]
        for s in speakers:
            lines.append(f"  \u2022 {s}")
        lines.append("")
        if self._session is not None and self._session.started_at is not None:
            lines.append(f"\u2705 Aktive Session: {self._session.title}")
            lines.append(f"   Gestartet:    {self._session.started_at.strftime('%H:%M:%S')}")
            lines.append(f"   Teilnehmer:   {len(self._session.participants)}")
        else:
            lines.append("\u26aa Keine aktive Meeting-Session.")

        QMessageBox.information(self, "Aktueller Status", "\n".join(lines))

    def _start_audio_pipeline(self) -> str:
        """Start the live audio capture + ASR buffering loop for the active meeting."""
        if self._active_meeting_id is None:
            return "Audio-Pipeline nicht gestartet (kein aktives Meeting)."

        self._clear_audio_buffer()
        self._asr_warned_no_text = False

        profile = self._selected_audio_profile()
        self._audio_capture_service = AudioCaptureService(profile=profile)
        self._mic_level_widget.set_initializing()
        try:
            self._audio_capture_service.start(self._on_audio_segment)
            self._asr_timer.start()
            self._mic_level_widget.set_active()
            return "Audioaufnahme aktiv, Live-Transkription läuft."
        except Exception as exc:
            logger.error("Audio-Pipeline konnte nicht gestartet werden: %s", exc)
            self._audio_capture_service = None
            self._mic_level_widget.set_error(str(exc))
            return f"Audio-Pipeline konnte nicht gestartet werden: {exc}"

    def _stop_audio_pipeline(self) -> None:
        self._asr_timer.stop()
        if self._audio_capture_service is not None:
            self._audio_capture_service.stop()
            self._audio_capture_service = None
        self._mic_level_widget.reset()

    def _clear_audio_buffer(self) -> None:
        with self._audio_buffer_lock:
            self._pending_audio_chunks = []
            self._pending_start_ms = None
            self._pending_end_ms = 0
            self._pending_duration_ms = 0

    def _on_audio_segment(self, segment: AudioSegment) -> None:
        """Collect non-silent chunks in a thread-safe buffer for periodic ASR."""
        # HEAR-044: feed the level meter (thread-safe via Signal/Slot)
        self._mic_level_widget.on_audio_segment(segment.rms, segment.is_silence)

        if self._active_meeting_id is None or segment.is_silence:
            return

        chunk = np.asarray(segment.samples, dtype=np.float32).reshape(-1)
        if chunk.size == 0:
            return

        with self._audio_buffer_lock:
            if self._pending_start_ms is None:
                self._pending_start_ms = segment.start_ms
            self._pending_audio_chunks.append(chunk)
            self._pending_end_ms = segment.end_ms
            self._pending_duration_ms += max(0, segment.end_ms - segment.start_ms)

    def _process_pending_audio(self) -> None:
        self._transcribe_pending_buffer(force=False)

    def _transcribe_pending_buffer(self, force: bool) -> None:
        payload = self._consume_audio_buffer(force=force)
        if payload is None:
            return

        start_ms, end_ms, samples = payload
        meeting_id = self._active_meeting_id
        if meeting_id is None:
            return

        segment = AudioSegment(
            captured_at=datetime.now(),
            start_ms=start_ms,
            end_ms=end_ms,
            samples=samples,
            rms=float(np.sqrt(np.mean(samples ** 2))),
            is_silence=False,
        )
        # ADR-0003: resolve speaker identity BEFORE persistence (HEAR-084 AC2)
        # Embedding-based pre-resolution ensures no segment is saved as unknown/0.0
        embedding = SpeakerManager._extract_embedding(samples.tolist())
        speaker_pre_match = self._speaker_manager.match_segment(embedding)

        result = self._transcription_service.transcribe_segment(
            segment,
            meeting_id=meeting_id,
            speaker_name=speaker_pre_match.speaker_name,
            confidence_score=speaker_pre_match.confidence,
        )

        text = result.text.strip() if result.text else ""
        # Refine with text hint (intro matching) — if improved, speaker_match overrides
        speaker_match = self._speaker_manager.resolve_speaker_from_segment(
            embedding, segment_text=text
        )

        stamp = self._format_ms(start_ms)
        review_tag = " [low-conf]" if speaker_match.requires_review else ""
        if text:
            self.transcript_line_ready.emit(
                f"[{stamp}] {speaker_match.speaker_name}{review_tag}: {text}"
            )
            return

        if (result.error or not self._asr_warned_no_text) and not text:
            self.transcript_line_ready.emit(
                f"[{stamp}] System: Audio erkannt, aber noch kein Transkriptions-Text verfügbar."
            )
            self._asr_warned_no_text = True

    def _consume_audio_buffer(self, force: bool) -> tuple[int, int, np.ndarray] | None:
        with self._audio_buffer_lock:
            if not self._pending_audio_chunks:
                return None

            if not force and self._pending_duration_ms < 1800:
                return None

            start_ms = self._pending_start_ms or 0
            end_ms = self._pending_end_ms
            chunks = self._pending_audio_chunks

            self._pending_audio_chunks = []
            self._pending_start_ms = None
            self._pending_end_ms = 0
            self._pending_duration_ms = 0

        return start_ms, end_ms, np.concatenate(chunks)

    @staticmethod
    def _format_ms(total_ms: int) -> str:
        total_seconds = max(0, total_ms // 1000)
        minutes = total_seconds // 60
        seconds = total_seconds % 60
        return f"{minutes:02d}:{seconds:02d}"

    def _build_transcript_panel(self) -> QWidget:
        panel = QWidget()
        layout = QVBoxLayout(panel)
        layout.setSpacing(12)

        status = QGroupBox("Live Transcript")
        status_layout = QVBoxLayout(status)
        status_layout.addWidget(QLabel("Input: default microphone | Whisper profile: balanced"))

        self._transcript_view = QPlainTextEdit()
        self._transcript_view.setReadOnly(True)
        self._transcript_view.setPlainText("[00:00] System: Meeting initialized.")
        status_layout.addWidget(self._transcript_view)
        layout.addWidget(status)
        return panel

    def _build_protocol_panel(self) -> QWidget:
        panel = QWidget()
        layout = QVBoxLayout(panel)
        layout.setSpacing(12)

        protocol_box = QGroupBox("Protocol Draft")
        protocol_layout = QVBoxLayout(protocol_box)
        protocol_layout.addWidget(QLabel(
            "Decisions, action items and open questions update on significant events."
        ))

        self._protocol_view = QPlainTextEdit()
        self._protocol_view.setReadOnly(True)
        self._protocol_view.setPlainText(
            "Summary\n"
            "- Meeting started with speaker enrollment.\n\n"
            "Decisions\n"
            "- Waiting for first confirmed decision.\n\n"
            "Action Items\n"
            "- No action items detected yet."
        )
        protocol_layout.addWidget(self._protocol_view)

        # HEAR-075: visible export action
        export_row = QHBoxLayout()
        self._export_btn = QPushButton("Export Protocol…")
        self._export_btn.setToolTip(
            "Protokoll als Markdown-Datei in den exports/-Ordner speichern"
        )
        self._export_btn.clicked.connect(self._do_export_protocol)
        self._export_btn.setEnabled(False)
        export_row.addWidget(self._export_btn)
        export_row.addStretch(1)
        protocol_layout.addLayout(export_row)

        self._export_path_label = QLabel("")
        self._export_path_label.setStyleSheet("color: #1a4a7a; font-size: 11px;")
        self._export_path_label.setWordWrap(True)
        protocol_layout.addWidget(self._export_path_label)

        layout.addWidget(protocol_box)

        quality_box = QGroupBox("Review Queue")
        quality_layout = QVBoxLayout(quality_box)
        quality_layout.addWidget(QLabel("Unknown speakers and low-confidence assignments appear here."))

        self._review_list = QListWidget()
        quality_layout.addWidget(self._review_list)

        self._speaker_override = QComboBox()
        self._speaker_override.setEditable(True)
        self._speaker_override.setPlaceholderText("Assign correct speaker name…")
        quality_layout.addWidget(self._speaker_override)

        correct_btn = QPushButton("Apply Correction")
        correct_btn.clicked.connect(self._apply_speaker_correction)
        quality_layout.addWidget(correct_btn)

        layout.addWidget(quality_box)
        return panel

    # ------------------------------------------------------------------
    # Review Queue: load + correction logic (HEAR-016)
    # ------------------------------------------------------------------

    def set_active_meeting(self, meeting_id: str, known_speakers: list[str] | None = None) -> None:
        """Called by orchestrator when a meeting session becomes active."""
        self._active_meeting_id = meeting_id

        self._speaker_override.clear()
        for name in (known_speakers or []):
            self._speaker_override.addItem(name)

        self._refresh_review_queue()
        self._refresh_timer.start()

    def stop_active_meeting(self) -> None:
        self._refresh_timer.stop()
        self._active_meeting_id = None
        self._speaker_manager.clear_meeting_context()

    def _refresh_review_queue(self) -> None:
        """Reload low-confidence and uncorrected segments from the repository."""
        self._review_list.clear()

        if self._transcript_repo is None or self._active_meeting_id is None:
            item = QListWidgetItem("(No active meeting or repository not connected)")
            item.setFlags(item.flags() & ~Qt.ItemFlag.ItemIsSelectable)
            self._review_list.addItem(item)
            return

        try:
            segments = self._transcript_repo.low_confidence(
                self._active_meeting_id,
                threshold=self.runtime_config.protocol.minimum_confidence,
            )
        except Exception as exc:
            logger.error("Failed to load review queue: %s", exc)
            return

        for seg in segments:
            label = (
                f"[{seg.start_ms}ms] {seg.speaker_name} "
                f"(confidence={seg.confidence_score:.2f}): {seg.text[:60]}"
            )
            item = QListWidgetItem(label)
            item.setData(Qt.ItemDataRole.UserRole, seg.id)
            self._review_list.addItem(item)

    def _apply_speaker_correction(self) -> None:
        """Persist a manual speaker correction for the selected review queue item."""
        selected = self._review_list.currentItem()
        corrected_name = self._speaker_override.currentText().strip()

        if selected is None or not corrected_name:
            QMessageBox.warning(self, "Correction", "Select a segment and enter a speaker name.")
            return

        segment_id: str = selected.data(Qt.ItemDataRole.UserRole)

        if self._transcript_repo is None:
            logger.warning("No transcript repository — correction not persisted.")
            return

        try:
            self._transcript_repo.apply_correction(segment_id, corrected_name)
            logger.info("Manual correction applied: segment=%s speaker=%s", segment_id, corrected_name)
        except Exception as exc:
            QMessageBox.critical(self, "Error", f"Could not save correction:\n{exc}")
            return

        # Remove from queue and refresh protocol snapshot if available
        row = self._review_list.row(selected)
        self._review_list.takeItem(row)
        self._refresh_protocol_display()

    def _refresh_protocol_display(self) -> None:
        """Reload the latest protocol snapshot and update the protocol view.

        Shows an explicit degraded label when persistence is not connected,
        so the panel clearly communicates its state per HEAR-085 AC3.
        """
        # AC3: explicit degraded label when persistence unavailable
        if self._snapshot_repo is None or self._active_meeting_id is None:
            current = self._protocol_view.toPlainText()
            if not current.startswith(_PROTOCOL_DEGRADED_PREFIX):
                self._protocol_view.setPlainText(
                    f"{_PROTOCOL_DEGRADED_PREFIX} Protokollentwurf nicht verf\u00fcgbar \u2014 "
                    "Persistenz-Backend nicht verbunden.\n\n"
                    "Verbinden Sie die Datenbank, um die strukturierte Protokollerstellung "
                    "gem\u00e4\u00df ADR-0005 zu aktivieren.\n\n"
                    "Live-Transkript siehe Transkript-Panel."
                )
            return
        try:
            snapshot = self._snapshot_repo.latest(self._active_meeting_id)
            if snapshot is None:
                return
            content = snapshot.snapshot_content or {}
            lines: list[str] = []

            def _section(title: str, items: list[str]) -> None:
                lines.append(title)
                for item in items:
                    lines.append(f"- {item}")
                lines.append("")

            _section("Summary", content.get("summary", []))
            _section("Decisions", content.get("decisions", []))
            _section("Action Items", content.get("action_items", []))
            _section("Open Questions", content.get("open_questions", []))

            self._protocol_view.setPlainText("\n".join(lines).strip())
        except Exception as exc:
            logger.error("Protocol refresh failed: %s", exc)

    def _rebuild_protocol_from_persistence(self) -> None:
        """Generate a new protocol snapshot from persisted transcript data (HEAR-085 AC2).

        Called by the periodic refresh timer (every 10 s).  Uses ProtocolEngine
        to produce a structured draft from confirmed transcript segments and
        stores it as a new versioned snapshot.  No-ops when persistence is
        unavailable or no meeting is active.
        """
        if (
            self._active_meeting_id is None
            or self._transcript_repo is None
            or self._snapshot_repo is None
        ):
            return
        try:
            self._protocol_engine.generate(self._active_meeting_id)
            self._refresh_protocol_display()
        except Exception as exc:
            logger.error("Protocol rebuild failed: %s", exc)

    # ------------------------------------------------------------------
    # HEAR-070 / HEAR-075: Protocol export (multi-format)
    # ------------------------------------------------------------------

    def _resolve_export_dir(self) -> "Path":
        """Return the directory where meeting artifacts are written (HEAR-073 / ADR-0011).

        Delegates to ``ayehear.utils.paths.exports_dir()`` so the path honours
        ``AYEHEAR_INSTALL_DIR`` and install-root-relative semantics.  Extracted
        as a dedicated method so tests can patch it without touching the
        filesystem helper.
        """
        from pathlib import Path as _Path
        from ayehear.utils.paths import exports_dir as _exports_dir
        return _exports_dir()

    def _export_meeting_artifacts(
        self,
        meeting_id: str | None,
        title: str,
    ) -> "list[Path]":
        """Export the current protocol and transcript as multi-format artifacts.

        Returns a list of ``Path`` objects for every file that was written.
        Returns an empty list when ``meeting_id`` is ``None`` or both the
        protocol and transcript views are empty.

        Protocol is written as Markdown, DOCX and PDF.
        Transcript is written as plain text (``-transcript.txt``).
        """
        import datetime as _dt
        from pathlib import Path as _Path

        if meeting_id is None:
            return []

        draft = self._protocol_view.toPlainText().strip()
        transcript = self._transcript_view.toPlainText().strip()

        if not draft and not transcript:
            return []

        meeting_type = ""
        if self._session is not None:
            meeting_type = self._meeting_type.currentText() if hasattr(self, "_meeting_type") else ""

        timestamp = _dt.datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_title = "".join(c if c.isalnum() or c in "-_" else "_" for c in title)[:40] or "meeting"
        base_name = f"{safe_title}_{timestamp}"

        try:
            out_dir: _Path = self._resolve_export_dir()
            out_dir.mkdir(parents=True, exist_ok=True)
        except OSError as exc:
            logger.error("Cannot create export directory: %s", exc)
            return []

        written: list[_Path] = []

        if draft:
            md_text = self._format_as_markdown(draft, title, meeting_type)
            # Markdown
            md_path = out_dir / f"{base_name}-protocol.md"
            try:
                md_path.write_text(md_text, encoding="utf-8")
                written.append(md_path)
            except OSError as exc:
                logger.error("Markdown export failed: %s", exc)

            # DOCX
            docx_path = out_dir / f"{base_name}-protocol.docx"
            try:
                from docx import Document as _Document  # type: ignore[import-untyped]
                doc = _Document()
                doc.add_heading(title, level=1)
                for line in md_text.splitlines():
                    if line.startswith("## "):
                        doc.add_heading(line[3:], level=2)
                    elif line.startswith("# "):
                        doc.add_heading(line[2:], level=1)
                    else:
                        doc.add_paragraph(line)
                doc.save(str(docx_path))
                written.append(docx_path)
            except Exception as exc:
                logger.error("DOCX export failed: %s", exc)

            # PDF via reportlab
            pdf_path = out_dir / f"{base_name}-protocol.pdf"
            try:
                from reportlab.lib.pagesizes import A4  # type: ignore[import-untyped]
                from reportlab.platypus import SimpleDocTemplate, Paragraph  # type: ignore[import-untyped]
                from reportlab.lib.styles import getSampleStyleSheet  # type: ignore[import-untyped]
                styles = getSampleStyleSheet()
                story = []
                for line in md_text.splitlines():
                    style = styles["Heading2"] if line.startswith("## ") else styles["Normal"]
                    text = line.lstrip("#").strip() or "\u00a0"
                    story.append(Paragraph(text, style))
                doc_pdf = SimpleDocTemplate(str(pdf_path), pagesize=A4)
                doc_pdf.build(story)
                written.append(pdf_path)
            except Exception as exc:
                logger.error("PDF export failed: %s", exc)

        # Transcript TXT
        if transcript:
            txt_path = out_dir / f"{base_name}-transcript.txt"
            try:
                header = f"Meeting Transcript — {title}\nExported: {_dt.datetime.now().isoformat()}\n\n"
                txt_path.write_text(header + transcript, encoding="utf-8")
                written.append(txt_path)
            except OSError as exc:
                logger.error("Transcript export failed: %s", exc)

        return written

    @staticmethod
    def _format_as_markdown(draft: str, title: str, meeting_type: str) -> str:
        """Convert a plain-text protocol draft to Markdown.

        Known section header names are converted to ``## Header`` lines.
        All other lines are preserved verbatim.
        """
        import datetime as _dt
        _SECTIONS = {"Summary", "Decisions", "Action Items", "Open Questions", "Transcript"}
        lines = [
            f"# {title}",
            "",
            f"**Type:** {meeting_type}",
            f"**Date:** {_dt.datetime.now().strftime('%Y-%m-%d %H:%M')}",
            "",
        ]
        for raw in draft.splitlines():
            stripped = raw.strip()
            if stripped in _SECTIONS:
                lines.append(f"## {stripped}")
            else:
                lines.append(raw)
        return "\n".join(lines)

    def _do_export_protocol(self) -> None:
        """Export the current protocol draft to <install_root>/exports/ as Markdown."""
        import datetime as _dt
        from ayehear.utils.paths import exports_dir as _exports_dir

        draft = self._protocol_view.toPlainText().strip()
        if not draft or draft.startswith("[DEGRADED]"):
            if draft.startswith("[DEGRADED]"):
                QMessageBox.warning(self, "Export", "Kein exportierbarer Protokollentwurf vorhanden.\n\nDas Protokoll ist degradiert — Datenbankverbindung erforderlich.")
            else:
                QMessageBox.warning(self, "Export", "Das Protokoll ist noch leer — nichts zu exportieren.")
            return

        title = ""
        if self._session is not None:
            title = self._session.title or ""

        timestamp = _dt.datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_title = "".join(c if c.isalnum() or c in "-_" else "_" for c in title)[:40]
        filename = f"protocol_{safe_title}_{timestamp}.md" if safe_title else f"protocol_{timestamp}.md"

        try:
            out_dir = _exports_dir()
            out_path = out_dir / filename
            header = f"# Meeting Protocol — {title}\n\nExported: {_dt.datetime.now().isoformat()}\n\n"
            out_path.write_text(header + draft, encoding="utf-8")
            self._export_path_label.setText(f"Exportiert: {out_path}")
            logger.info("Protocol exported to %s", out_path)
            QMessageBox.information(
                self,
                "Export erfolgreich",
                f"Protokoll gespeichert:\n{out_path}",
            )
        except OSError as exc:
            logger.error("Protocol export failed: %s", exc)
            QMessageBox.critical(self, "Export fehlgeschlagen", f"Fehler beim Schreiben:\n{exc}")

    def _update_protocol_live(self, transcript_line: str) -> None:  # noqa: ARG002
        """Refresh the protocol draft view on each new transcript line (HEAR-085 AC1/AC2).

        The protocol panel represents a *structured draft*, never a transcript
        mirror.  When persistence is available, delegates to the snapshot-based
        refresh.  When persistence is unavailable, the degraded-state label
        remains visible — transcript lines are ONLY added to the transcript view,
        never to the protocol panel.
        """
        if self._snapshot_repo is not None and self._active_meeting_id is not None:
            # With DB: refresh the structured snapshot view
            self._refresh_protocol_display()
            return
        # Without DB: ensure degraded label is visible (do NOT mirror transcript)
        current = self._protocol_view.toPlainText()
        if not current.startswith(_PROTOCOL_DEGRADED_PREFIX):
            self._protocol_view.setPlainText(
                f"{_PROTOCOL_DEGRADED_PREFIX} Protokollentwurf nicht verf\u00fcgbar \u2014 "
                "Persistenz-Backend nicht verbunden.\n\n"
                "Live-Transkript siehe Transkript-Panel."
            )

    def append_transcript_line(self, line: str) -> None:
        """Append a new transcribed line to the live transcript view (HEAR-075: also updates protocol)."""
        self._transcript_view.appendPlainText(line)
        # Scroll to bottom
        self._transcript_view.verticalScrollBar().setValue(
            self._transcript_view.verticalScrollBar().maximum()
        )
        # HEAR-075: keep protocol draft visible and up-to-date during meeting
        if self._active_meeting_id is not None:
            self._update_protocol_live(line)
